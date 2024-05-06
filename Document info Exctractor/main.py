import streamlit as st
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import Workbook


def extract_info_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    for page_num in range(len(reader.pages)):
        text += reader.pages[page_num].extract_text()  # Use extract_text() instead of extract_Text()
    doc_number = text.split("Document Number:")[0].split("\n")[0].strip()
    title = text.split("Title:")[0].split("\n")[0].strip()
    revision = text.split("Revision:")[0].split("\n")[0].strip()
    return doc_number, title, revision

def extract_info_docx(docx_file):
    doc = Document(docx_file)
    doc_number = doc.paragraphs[0].text.split(":")[1].strip()
    title = doc.paragraphs[1].text.split(":")[1].strip()
    revision = doc.paragraphs[2].text.split(":")[1].strip()
    return doc_number, title, revision

def main():
    st.title("Document Info Extractor")
    st.write("Upload a PDF or Word document to extract document information.")

    uploaded_file = st.file_uploader("Choose a file...", type=["pdf", "docx"])

    if uploaded_file is not None:
        file_type = uploaded_file.name.split(".")[-1]
        if file_type == "pdf":
            doc_number, title, revision = extract_info_pdf(uploaded_file)
        elif file_type == "docx":
            doc_number, title, revision = extract_info_docx(uploaded_file)
        else:
            st.error("Unsupported file format. Please upload a PDF or Word document.")
            return

        st.write("Document Number:", doc_number)
        st.write("Title:", title)
        st.write("Revision:", revision)

        if st.button("Save to Excel"):
            df = pd.DataFrame({
                "Document Number": [doc_number],
                "Title": [title],
                "Revision": [revision]
            })
            df.to_excel("document_info.xlsx", index=False)
            st.success("Document information saved to Excel file.")

if __name__ == "__main__":
    main()
